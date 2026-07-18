from __future__ import annotations

from dataclasses import dataclass
from io import BytesIO
from pathlib import Path

import fitz
from docx import Document

from src.config.settings import settings
from src.utils.text import clean_text


class DocumentExtractionError(ValueError):
    pass


@dataclass(frozen=True)
class ExtractedDocument:
    file_name: str
    text: str
    page_count: int | None = None
    warning: str | None = None


class DocumentExtractor:
    supported_extensions = {".pdf", ".docx", ".txt"}

    def extract(self, file_name: str, data: bytes) -> ExtractedDocument:
        suffix = Path(file_name).suffix.lower()
        if suffix not in self.supported_extensions:
            raise DocumentExtractionError(f"Unsupported file type: {suffix or 'unknown'}")
        if not data:
            raise DocumentExtractionError(f"{file_name} is empty.")
        if len(data) > settings.max_file_size_mb * 1024 * 1024:
            raise DocumentExtractionError(
                f"{file_name} exceeds the {settings.max_file_size_mb} MB limit."
            )

        try:
            if suffix == ".pdf":
                result = self._extract_pdf(file_name, data)
            elif suffix == ".docx":
                result = self._extract_docx(file_name, data)
            else:
                result = ExtractedDocument(file_name, clean_text(data.decode("utf-8", errors="ignore")))
        except DocumentExtractionError:
            raise
        except Exception as exc:
            raise DocumentExtractionError(f"Could not parse {file_name}: {exc}") from exc

        if len(result.text) < settings.min_extracted_characters:
            raise DocumentExtractionError(
                f"Very little text was extracted from {file_name}. It may be scanned, damaged, or empty."
            )
        return result

    @staticmethod
    def _extract_pdf(file_name: str, data: bytes) -> ExtractedDocument:
        doc = fitz.open(stream=data, filetype="pdf")
        pages = [page.get_text("text", sort=True) for page in doc]
        text = clean_text("\n".join(pages))
        warning = None
        if len(text) < 200:
            warning = "Low text density detected; this may be an image-only PDF requiring OCR."
        return ExtractedDocument(file_name, text, len(doc), warning)

    @staticmethod
    def _extract_docx(file_name: str, data: bytes) -> ExtractedDocument:
        doc = Document(BytesIO(data))
        blocks: list[str] = [p.text for p in doc.paragraphs if p.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                blocks.append(" | ".join(cell.text.strip() for cell in row.cells))
        return ExtractedDocument(file_name, clean_text("\n".join(blocks)))
